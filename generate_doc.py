from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────

def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(d, text, level=1):
    p = d.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    colors = {1: (0x1a,0x56,0xdb), 2: (0x1e,0x40,0xaf), 3: (0x37,0x51,0x8a)}
    sizes  = {1: 18, 2: 14, 3: 12}
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor(*colors.get(level, (0,0,0)))
    return p

def add_para(d, text, bold=False, italic=False, size=11, color=None):
    p = d.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(d, text):
    p = d.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(d, headers, rows, col_widths=None):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style     = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, bold=True, size=10, color=(0xff,0xff,0xff))
        shade_cell(cell, '1e40af')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for ri, row_data in enumerate(rows):
        fill = 'f0f4ff' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
            shade_cell(cell, fill)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    d.add_paragraph()
    return t

def add_code(d, text):
    for line in text.strip().split('\n'):
        p = d.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
    d.add_paragraph()

def add_divider(d):
    p   = d.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '93c5fd')
    pBdr.append(bot)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(3)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('RIDER ASSET MANAGEMENT SYSTEM')
run.font.name  = 'Calibri'
run.font.size  = Pt(26)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1a,0x56,0xdb)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run('Engineering Manager Briefing Document')
run2.font.name    = 'Calibri'
run2.font.size    = Pt(14)
run2.font.italic  = True
run2.font.color.rgb = RGBColor(0x64,0x74,0x8b)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
for label, value in [
    ('Date',       'March 22, 2026'),
    ('Author',     'Devansh'),
    ('Repository', 'github.com/Devanshyadv/Rider_Asset_management'),
    ('Status',     'Complete — Deployed & Tested'),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(label + ':  ')
    set_font(r1, bold=True, size=11, color=(0x1e,0x40,0xaf))
    r2 = p.add_run(value)
    set_font(r2, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════
add_heading(doc, '1. Executive Summary', 1)
add_divider(doc)
add_para(doc,
    'This system automates the full lifecycle of physical asset management for last-mile delivery riders — '
    'from issuance and cost recovery through weekly payroll deductions, to vendor liability reconciliation '
    'and Finance-team override controls.')
add_para(doc,
    'Before this system, cost recovery was manual — spreadsheet-driven, error-prone, and lacked any audit trail. '
    'This backend eliminates that entirely.')

add_heading(doc, 'What It Delivers', 3)
for b in [
    'Issues assets to riders under FLAT (single deduction) or EMI (up to 4-cycle) repayment plans',
    'Runs a weekly settlement engine every Monday, deducting from rider earnings in a strict priority order',
    'Tracks vendor liabilities and records FIFO vendor payout settlements',
    'Provides Finance and LMD teams admin override controls: waive, adjust, flag abuse, year-end write-off',
    'Logs every financial action to an immutable audit trail',
    'Ships a browser-based admin dashboard — no toolchain required',
]:
    add_bullet(doc, b)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  2. BUSINESS PROBLEMS SOLVED
# ══════════════════════════════════════════════════════════
add_heading(doc, '2. Business Problems Solved', 1)
add_divider(doc)
add_table(doc,
    ['Pain Point', 'How the System Addresses It'],
    [
        ('No central asset catalog',        'Asset master with stock management and soft-delete'),
        ('Manual deduction tracking',       'Automated weekly settlement engine with carry-forwards'),
        ('Unclear vendor accountability',   'Per-vendor liability ledger with FIFO payout settlement'),
        ('No Finance override capability',  'Waive, adjust deductions, flag abuse, year-end write-off'),
        ('No audit trail',                  'Every state change writes an immutable log entry'),
        ('No rider visibility',             'Rider-facing API shows active assets and outstanding balance'),
        ('Slow bulk onboarding',            'CSV-based bulk issuance upload with per-row error reporting'),
    ],
    col_widths=[7, 10]
)

# ══════════════════════════════════════════════════════════
#  3. DELIVERED SCOPE
# ══════════════════════════════════════════════════════════
add_heading(doc, '3. Delivered Scope', 1)
add_divider(doc)

add_heading(doc, '3.1  Backend API — 20 Endpoints across 7 Domains', 2)
add_table(doc,
    ['Domain', 'Endpoints', 'Purpose'],
    [
        ('Assets',     '4', 'Create, list, update, soft-delete asset catalog'),
        ('Issuances',  '4', 'Issue asset, get by rider, bulk CSV upload, replace'),
        ('Settlement', '2', 'Run weekly engine, check status by week'),
        ('Vendors',    '2', 'View liabilities, record payment'),
        ('Reports',    '3', 'Settlement report, vendor liability summary, audit trail'),
        ('Admin',      '4', 'Waive, adjust deduction, flag abuse, year-end write-off'),
        ('Rider View', '1', "Rider's active assets and outstanding balances"),
    ],
    col_widths=[4, 3, 11]
)

add_heading(doc, '3.2  Admin Frontend — 9 Sections', 2)
add_para(doc,
    'Dashboard  ·  Asset Catalog  ·  Issue Asset  ·  Bulk Upload  ·  Rider Lookup  ·  '
    'Settlement  ·  Vendor Management  ·  Reports  ·  Admin Actions',
    italic=True)
add_para(doc, 'Opened directly in any browser — no server, no build step, no dependencies.')

add_heading(doc, '3.3  Test Suite', 2)
add_table(doc,
    ['Test File', 'Covers'],
    [
        ('test_api_endpoints.py (77 tests)',   'All 20 endpoints — happy paths, validation errors, all edge cases'),
        ('test_issuance_validation.py',         'FLAT/EMI rules, cycle cap, duplicate logic'),
        ('test_settlement_engine.py',           'Priority ordering, carry-forward, vendor pass-through'),
    ],
    col_widths=[6, 11]
)
add_para(doc,
    'Full HTTP-level testing via httpx AsyncClient. Zero live database dependency. Runs in ~3-4 seconds.',
    italic=True, color=(0x64,0x74,0x8b))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  4. ARCHITECTURE
# ══════════════════════════════════════════════════════════
add_heading(doc, '4. Technical Architecture', 1)
add_divider(doc)

add_heading(doc, '4.1  Technology Stack', 2)
add_table(doc,
    ['Layer', 'Technology', 'Reason for Choice'],
    [
        ('API Framework',  'FastAPI (Python 3.12)',   'Async-native, auto-generates OpenAPI docs, Pydantic validation built-in'),
        ('Database',       'MongoDB Atlas',           'Flexible document model for varied asset/issuance data; free tier available'),
        ('DB Driver',      'Motor (async)',           'Non-blocking I/O — essential for concurrent settlement runs'),
        ('Validation',     'Pydantic v2',             'Strict schema enforcement at every API boundary, zero boilerplate'),
        ('Scheduler',      'APScheduler',             'Triggers weekly settlement every Monday 00:00 UTC automatically'),
        ('CSV Processing', 'Pandas',                  'Robust bulk upload with per-row error reporting'),
        ('Testing',        'pytest + httpx',          'Full async test support; no live server or database needed'),
        ('Frontend',       'Vanilla JS / HTML / CSS', 'Zero build toolchain — open index.html in any browser'),
    ],
    col_widths=[3.5, 4, 9.5]
)

add_heading(doc, '4.2  Key Engineering Decisions', 2)
add_table(doc,
    ['Decision', 'Rationale'],
    [
        ('Snapshot fields on issuances',
         'costPriceSnapshot, absorptionTypeSnapshot, vendorIdSnapshot frozen at issue time. '
         'Changes to the asset master never retroactively affect active deduction plans.'),
        ('Soft deletes on assets',
         'Assets are never hard-deleted. status=inactive preserves referential integrity for all historical issuances.'),
        ('Dependency injection for DB',
         'get_db() injected via FastAPI Depends(). Makes the entire test suite mockable without any live database.'),
        ('Fully async codebase',
         'Every DB call uses await via Motor. Settlement runs never block the event loop.'),
    ],
    col_widths=[5, 12]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. SETTLEMENT ENGINE
# ══════════════════════════════════════════════════════════
add_heading(doc, '5. Settlement Engine — Core Business Logic', 1)
add_divider(doc)
add_para(doc,
    'The settlement engine is the most critical component. It runs automatically every Monday '
    'and can also be triggered manually via the API.')

add_heading(doc, '5.1  Deduction Priority Order (deterministic, non-negotiable)', 2)
add_table(doc,
    ['Priority', 'Type', 'Order Within'],
    [
        ('P0',  'SELF carry-forwards (unpaid shortfalls from prior weeks)', 'Oldest unpaid week first'),
        ('P1a', 'SELF active issuances (rider pays)',                       'Oldest issued first'),
        ('P1b', 'VENDOR active issuances (vendor absorbs)',                 'Oldest issued first'),
    ],
    col_widths=[2, 9, 6]
)
add_para(doc,
    'SELF debts are always settled before VENDOR debts. '
    'Prior-week carry-forwards are cleared before any new deductions.',
    italic=True, color=(0x1e,0x40,0xaf))

add_heading(doc, '5.2  Settlement Outcomes Per Deduction', 2)
add_table(doc,
    ['Outcome', 'Condition', 'Effect'],
    [
        ('deducted',      'Earnings cover full amount due',    'outstandingBalance reduced; cycle counted'),
        ('partial',       'Earnings partially cover amount',   'Shortfall written to rider_carry_forwards'),
        ('skipped',       'Earnings = 0 or fully exhausted',   'Full amount written to rider_carry_forwards'),
        ('completed',     'outstandingBalance reaches 0',      'Issuance status set to completed'),
        ('vendor_passed', 'VENDOR balance remains at end',     'Liability record created in vendor_liabilities'),
    ],
    col_widths=[3, 6, 8]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  6. ADMIN CONTROLS
# ══════════════════════════════════════════════════════════
add_heading(doc, '6. Admin Controls', 1)
add_divider(doc)
add_table(doc,
    ['Action', 'Who', 'What It Does'],
    [
        ('Waive Issuance',
         'Finance',
         'Forgives entire outstanding balance. Status set to written_off. All pending carry-forwards marked recovered.'),
        ('Adjust Deduction',
         'Finance',
         'Manually corrects amount on a ledger entry. Auto-reconciles totalRecovered and outstandingBalance on parent issuance.'),
        ('Flag Abuse',
         'LMD',
         'Overrides absorption from VENDOR to SELF. Rider becomes responsible for remaining balance. Reactivates vendor_passed issuances.'),
        ('Year-End Write-Off',
         'Finance',
         'Bulk writes off SELF active issuances where rider has no orders since financial year-end. Supports dry_run=true to preview first.'),
    ],
    col_widths=[4, 2.5, 10.5]
)
add_para(doc,
    'Every admin action records the performing user ID, reason, and before/after state in the audit log.',
    italic=True, color=(0x64,0x74,0x8b))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  7. API QUICK REFERENCE
# ══════════════════════════════════════════════════════════
add_heading(doc, '7. API Surface — Quick Reference', 1)
add_divider(doc)
add_code(doc, """GET  / POST        /api/assets
PATCH / DELETE     /api/assets/{id}

POST               /api/issuances               ?confirm=true
GET                /api/issuances/{rider_id}     ?status=active
POST               /api/issuances/bulk           (CSV multipart upload)
POST               /api/issuances/{id}/replace   ?mid_balance_action=waive_remaining|pass_to_vendor

POST               /api/settlement/run-weekly
GET                /api/settlement/status/{week}    e.g. 2026-W12

GET                /api/vendors/{vendor_id}/liabilities
POST               /api/vendors/{vendor_id}/settle

GET                /api/reports/settlement/{week}
GET                /api/reports/vendor-liability
GET                /api/reports/audit/{entity_id}

POST               /api/admin/waive/{issuance_id}
PATCH              /api/admin/deduction/{ledger_id}
POST               /api/admin/flag-abuse/{issuance_id}
POST               /api/admin/year-end-writeoff

GET                /api/rider/{rider_id}/assets""")
add_para(doc,
    'Interactive docs available at /docs (Swagger UI) and /redoc when server is running.',
    italic=True, color=(0x64,0x74,0x8b))
doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  8. DATABASE COLLECTIONS
# ══════════════════════════════════════════════════════════
add_heading(doc, '8. Database Collections Summary', 1)
add_divider(doc)
add_table(doc,
    ['Collection', 'Purpose', 'Key Fields'],
    [
        ('asset_master',           'Asset catalog',                    'name, costPrice, absorptionType, stock, status'),
        ('asset_issuances',        'All active and historical records', 'riderId, assetId, reconciliationType, totalCycles, outstandingBalance, status'),
        ('asset_deduction_ledger', 'One record per deduction cycle',   'weekCycle, amountDue, amountDeducted, shortfall, status'),
        ('vendor_liabilities',     'What each vendor owes',            'vendorId, outstandingAmount, amountRecovered, status'),
        ('rider_carry_forwards',   'Unpaid shortfalls rolled over',    'weekCycle, outstandingCarryForward, status'),
        ('asset_audit_logs',       'Immutable history of every action','action, before_value, after_value, reason, performed_by'),
        ('orders',                 'Rider order history',              'rider, createdAt  (used for year-end write-off eligibility check)'),
    ],
    col_widths=[4.5, 5, 7.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  9. KNOWN RISKS & GAPS
# ══════════════════════════════════════════════════════════
add_heading(doc, '9. Known Risks & Gaps', 1)
add_divider(doc)
add_table(doc,
    ['Risk', 'Severity', 'Status / Recommendation'],
    [
        ('No authentication — all endpoints are open',         'HIGH',   'JWT/OAuth2 required before any production exposure'),
        ('No rate limiting on settlement endpoint',            'MEDIUM', 'Can be triggered repeatedly; add slowapi or Nginx throttle'),
        ('CORS accepts all origins',                           'MEDIUM', 'Lock allow_origins to specific frontend domain'),
        ('No pagination on list endpoints',                    'LOW',    'Will degrade at scale; add limit/cursor pagination'),
        ('Settlement not batched for large rider counts',      'LOW',    'For >10,000 riders, introduce batching in SettlementEngine'),
        ('Bulk CSV upload is sequential',                      'LOW',    'Large files (>1,000 rows) will be slow; consider async chunking'),
    ],
    col_widths=[7, 2.5, 7.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  10. PRODUCTION HARDENING
# ══════════════════════════════════════════════════════════
add_heading(doc, '10. Recommended Production Hardening', 1)
add_divider(doc)
add_table(doc,
    ['Priority', 'Action'],
    [
        ('P0', 'Add JWT authentication with Finance / LMD / Rider roles. Gate /api/admin/* behind Finance role.'),
        ('P0', 'Create 7 MongoDB indexes across all collections for settlement and reporting query performance.'),
        ('P0', 'Set CORS allow_origins to the actual admin frontend domain only.'),
        ('P0', 'Move MONGODB_URL to AWS Secrets Manager or HashiCorp Vault.'),
        ('P1', 'Add rate limiting on /api/settlement/run-weekly.'),
        ('P1', 'Add limit/cursor pagination to list endpoints.'),
        ('P2', 'Set up MongoDB Atlas alerting on collection size thresholds.'),
        ('P2', 'Run as systemd service behind Nginx with SSL. Docker image included.'),
    ],
    col_widths=[2, 15]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  11. ROADMAP
# ══════════════════════════════════════════════════════════
add_heading(doc, '11. Roadmap', 1)
add_divider(doc)
add_table(doc,
    ['Priority', 'Item'],
    [
        ('P0', 'Authentication & RBAC (JWT with role-based access)'),
        ('P0', 'MongoDB production indexes'),
        ('P1', 'Pagination on all list endpoints'),
        ('P1', 'Rate limiting on settlement trigger'),
        ('P2', 'Settlement dry-run mode (preview before committing)'),
        ('P2', 'Rider mobile API (read-only stripped-down endpoints)'),
        ('P3', 'Email / SMS deduction notifications to riders'),
        ('P3', 'Analytics dashboard — settlement trends, vendor reconciliation charts'),
    ],
    col_widths=[2, 15]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  12. QUICK REFERENCE
# ══════════════════════════════════════════════════════════
add_heading(doc, '12. Quick Reference', 1)
add_divider(doc)
add_table(doc,
    ['Item', 'Value'],
    [
        ('GitHub Repository', 'https://github.com/Devanshyadv/Rider_Asset_management'),
        ('API Docs (local)',   'http://localhost:8000/docs'),
        ('Frontend',          'Open frontend/index.html in any browser'),
        ('Start server',      'uvicorn main:app --reload --port 8000'),
        ('Run tests',         'cd rider-asset-management && pytest -v'),
        ('Environment config','rider-asset-management/.env  (not committed to repo)'),
        ('Total endpoints',   '20 across 7 domains'),
        ('Total tests',       '77 HTTP-level integration tests'),
        ('Test runtime',      '~3-4 seconds (fully mocked, no network required)'),
    ],
    col_widths=[4.5, 12.5]
)

add_para(doc, '')
add_para(doc,
    'All items listed under Section 3 (Delivered Scope) are implemented, tested, and pushed to the repository.',
    italic=True, bold=True, color=(0x1e,0x40,0xaf))

# ── Save ──────────────────────────────────────────────────
out = r'C:\Assignment_fairdeal\EM_Briefing_RiderAssetManagement.docx'
doc.save(out)
print('Saved:', out)
